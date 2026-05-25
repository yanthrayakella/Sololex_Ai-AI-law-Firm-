"""Judge demo: merge a fake regulatory block into a work permit file (txt/docx/pdf)."""

import io
import re
import uuid
from pathlib import Path

from docx import Document
from pypdf import PdfMerger
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from config import DATA_DIR

DEMO_DIR = DATA_DIR / "demo_sessions"
ALLOWED = {".txt", ".docx", ".pdf"}

FAKE_REGULATION_TITLE = "NEW REGULATION DETECTED — March 28, 2026"

FAKE_REGULATION_BODY = """ADMINISTRATIVE NOTICE (DEMO — Simulated government website extract)
Source: Sample National Immigration / HR portal — NOT A REAL LEGAL INSTRUMENT

Effective March 28, 2026, the following items MUST be reflected on or appended to foreign work permits:

1. Verified in-country emergency contact
Registered through the official portal within 10 business days of this notice. Contact name, relationship, telephone, and city of residence required.

2. Digital Compliance Briefing (DCB-2026)
Permit holder must complete the online briefing, save the confirmation code, and keep proof available for inspection.

3. Compliance QR reference
The permit package must include the appendix line: "QR Ref: DCB-2026-CONFIRMED" where indicated by the issuing workflow.

This notice is fabricated for judge demonstration purposes only."""

SAMPLE_WORK_PERMIT = """================================================================================
                         WORK PERMIT — SAMPLE (Judge Demo)
================================================================================

Permit No.:          DEMO-WP-2026-0142
Full name:          ALEXANDER DEMO
Nationality:        Exampleland
Date of birth:      1990-05-12
Passport:           EP 1234567

Employer (Chinese): Shanghai Demo Technology Co., Ltd.
Unified credit code: 91310000MA00DEMO01
Job title:          Senior Software Engineer
Work location:      Shanghai Municipality

Valid from:         2026-01-10
Valid until:        2027-01-09

This SAMPLE permit authorizes the named holder to perform the stated employment
activities in accordance with applicable laws. This document is for demonstration
only and has no legal effect.

Issuing authority (demo): Municipal Human Resources & Social Security Bureau

Holder signature: _________________________    Date: _______________
Official stamp area: [ DEMO ]
"""


def ensure_demo_dir(session_id: str) -> Path:
    if not re.fullmatch(r"^[0-9a-f-]{36}$", session_id, re.I):
        raise ValueError("Invalid session")
    d = DEMO_DIR / session_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def find_original(path: Path) -> Path:
    for name in ("original.pdf", "original.docx", "original.txt"):
        p = path / name
        if p.is_file():
            return p
    raise FileNotFoundError("Original upload not found")


def _appendix_pdf_bytes(title: str, body: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 48
    c.setFont("Helvetica-Bold", 12)
    for line in title.split("\n"):
        c.drawString(48, y, line[:95])
        y -= 16
    y -= 8
    c.setFont("Helvetica", 10)
    for para in body.split("\n"):
        if not para.strip():
            y -= 10
            continue
        words = para.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, "Helvetica", 10) < width - 96:
                line = test
            else:
                c.drawString(48, y, line[:120])
                y -= 12
                line = w
            if y < 72:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 48
        if line:
            c.drawString(48, y, line[:120])
            y -= 12
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 48
    c.save()
    buf.seek(0)
    return buf.read()


def merge_pdf(original: Path, dest: Path, title: str, body: str) -> None:
    appendix_pdf = io.BytesIO(_appendix_pdf_bytes(title, body))
    appendix_pdf.seek(0)
    merger = PdfMerger()
    merger.append(str(original))
    merger.append(appendix_pdf)
    merger.write(str(dest))
    merger.close()


def merge_docx(original: Path, dest: Path, title: str, body: str) -> None:
    doc = Document(str(original))
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    for para in body.split("\n"):
        doc.add_paragraph(para)
    doc.save(str(dest))


def merge_txt(original: Path, dest: Path, title: str, body: str) -> None:
    text = original.read_text(encoding="utf-8", errors="ignore")
    block = f"\n\n{'=' * 60}\n{title}\n{'=' * 60}\n{body}\n"
    dest.write_text(text + block, encoding="utf-8")


def apply_fake_update(session_id: str) -> tuple[Path, str]:
    d = ensure_demo_dir(session_id)
    orig = find_original(d)
    suffix = orig.suffix.lower()
    out_name = f"updated_work_permit{suffix}"
    dest = d / out_name

    title = FAKE_REGULATION_TITLE
    body = FAKE_REGULATION_BODY

    if suffix == ".pdf":
        merge_pdf(orig, dest, title, body)
    elif suffix == ".docx":
        merge_docx(orig, dest, title, body)
    else:
        merge_txt(orig, dest, title, body)

    return dest, out_name


def save_upload(session_id: str, content: bytes, suffix: str) -> None:
    if suffix not in ALLOWED:
        raise ValueError("Invalid file type")
    d = ensure_demo_dir(session_id)
    dest = d / f"original{suffix}"
    dest.write_bytes(content)
